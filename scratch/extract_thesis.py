import docx
import os

def extract_hu_info(file_path, target_hu="HU010"):
    if not os.path.exists(file_path):
        print(f"File not found: {file_path}")
        return

    doc = docx.Document(file_path)
    content = []
    
    print(f"Searching for {target_hu} in paragraphs...")
    for para in doc.paragraphs:
        if target_hu in para.text:
            content.append(f"[Para] {para.text.strip()}")
    
    print(f"Searching for {target_hu} in tables...")
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(target_hu in cell_text for cell_text in row_text):
                content.append(f"[Table Row] {' | '.join(row_text)}")
                
    if content:
        print("\n".join(content))
    else:
        print(f"{target_hu} not found in the document.")

if __name__ == "__main__":
    path = r"c:\CerebroVial\documentation\tesis\TB1-251-223-03-BH-u20241c919-u202418685 (1).docx"
    extract_hu_info(path)
